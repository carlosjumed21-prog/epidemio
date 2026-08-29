import streamlit as st
import pandas as pd
import numpy as np
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuración de la página
st.set_page_config(
    page_title="Sistema de Evaluación Epidemiológica - SUIVE",
    page_icon="📊",
    layout="wide"
)

# Estilos CSS personalizados para la interfaz web
st.markdown("""
<style>
    .main-header { font-size: 2.2rem; color: #1E3A8A; font-weight: 700; margin-bottom: 0.2rem; }
    .sub-header { font-size: 1.1rem; color: #4B5563; margin-bottom: 1.5rem; }
    .info-box { background-color: #F8FAFC; border-left: 4px solid #1E3A8A; padding: 12px; margin-bottom: 20px; border-radius: 4px; }
    .legend-container { display: flex; gap: 15px; margin-bottom: 20px; flex-wrap: wrap; }
    .legend-item { padding: 8px 15px; border-radius: 6px; font-weight: bold; font-size: 0.9rem; text-align: center; }
    .legend-excelente { background-color: #10B981; color: white; }
    .legend-bueno { background-color: #FFFFFF; color: black; border: 1px solid #CBD5E1; }
    .legend-regular { background-color: #FEF08A; color: black; }
    .legend-malo { background-color: #EF4444; color: white; }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header">Evaluación de Indicadores Epidemiológicos SUAVE / SUIVE</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Módulo de Análisis por Trimestre, Semaforización y Reporte Word</div>', unsafe_allow_html=True)

# Lista completa de las 16 unidades operativas oficiales
TARGET_UNITS = [
    "CHURUBUSCO", "CLIDDA", "CMN 20 DE NOVIEMBRE", "COYOACAN", "DEL VALLE", 
    "DIVISION DEL NORTE", "DR. DARIO FERNANDEZ FIERRO", "DR. IGNACIO CHAVEZ", "ERMITA",
    "FUENTES BROTANTES", "HG DRA. MATILDE PETRA MONTOYA LAFRAGUA",
    "MILPA ALTA", "NARVARTE", "TLALPAN", "VILLA ALVARO OBREGON", "XOCHIMILCO"
]

# Subir archivo Excel
uploaded_file = st.file_uploader("📂 Sube tu archivo Excel de reportes SUIVE", type=["xlsx", "xls"])

if uploaded_file is not None:
    try:
        df = pd.read_excel(uploaded_file, sheet_name=0, header=None)
        
        delegacion = df.iloc[0, 1] if df.shape[0] > 0 and df.shape[1] > 1 else "REPRESENTACIÓN REGIONAL SUR"
        anio = int(df.iloc[1, 1]) if df.shape[0] > 1 and df.shape[1] > 1 and str(df.iloc[1, 1]).isdigit() else 2024
        
        # Detección de Año Bisiesto
        es_bisiesto = (anio % 4 == 0 and anio % 100 != 0) or (anio % 400 == 0)
        total_semanas_anio = 53 if es_bisiesto else 52

        semanas_list = []
        if df.shape[0] > 4:
            for col_idx in range(1, 27):
                val_sem = df.iloc[4, col_idx]
                if pd.notna(val_sem):
                    semanas_list.append(str(val_sem).strip())
        
        total_semanas_reportadas = len(semanas_list)
        periodo_str = f"Semana {semanas_list[0]} a Semana {semanas_list[-1]} (Total: {total_semanas_reportadas} semanas)" if semanas_list else "No determinado"

        st.markdown(f"""
        <div class="info-box">
            <h4>📋 Información General y Estructura Trimestral</h4>
            <ul>
                <li><b>Delegación:</b> {delegacion}</li>
                <li><b>Año Analizado:</b> {anio} {('(Bisiesto - 53 semanas)' if es_bisiesto else '(No Bisiesto - 52 semanas)')}</li>
                <li><b>Estandarización Trimestral:</b> 13 semanas por periodo</li>
                <li><b>Periodo Detectado en Archivo:</b> {periodo_str}</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)

        # Mapeo y extracción de datos por unidad desde el Excel
        data_dict = {}
        current_unit = None
        metrics = {}
        
        for i, row in df.iterrows():
            val = row[0]
            ab_val = row[27] if len(row) > 27 else np.nan
            
            if pd.notna(val):
                val_str = str(val).strip()
                if val_str in TARGET_UNITS:
                    if current_unit:
                        data_dict[current_unit] = metrics
                    current_unit = val_str
                    metrics = {}
                elif current_unit:
                    if val_str in [
                        "Casos acumulados", "Casos oportunos", "Semanas acumuladas con casos",
                        "Unidades con casos oportunos", "Unidades habilitadas", "Unidades sin notificar"
                    ]:
                        metrics[val_str] = float(ab_val) if pd.notna(ab_val) else 0.0

        if current_unit:
            data_dict[current_unit] = metrics

        if not data_dict:
            st.error("No se encontraron unidades válidas en la Columna A con los nombres esperados.")
        else:
            st.success(f"¡Archivo procesado con éxito! Se mapearon correctamente las unidades.")
            
            # -------------------------------------------------------------
            # SELECTOR DE TRIMESTRE EN STREAMLIT
            # -------------------------------------------------------------
            st.markdown("---")
            trimestre_seleccionado = st.selectbox(
                "📅 Seleccione el Trimestre de Evaluación:", 
                ["PRIMER TRIMESTRE (Sem. 1-13)", "SEGUNDO TRIMESTRE (Sem. 14-26)", "TERCER TRIMESTRE (Sem. 27-39)", "CUARTO TRIMESTRE (Sem. 40-52/53)"]
            )

            # Cada trimestre normado equivale a 13 semanas
            divisor_periodo = 13.0

            processed_results = []
            for unidad in TARGET_UNITS:
                m = data_dict.get(unidad, {})
                semanas_casos = m.get("Semanas acumuladas con casos", 0.0)
                u_oportunas = m.get("Unidades con casos oportunos", 0.0)
                u_habilitadas = m.get("Unidades habilitadas", 16.0)
                u_sin_notificar = m.get("Unidades sin notificar", 0.0)
                
                base_hab = u_habilitadas if u_habilitadas > 0 else 16.0

                promedio_semanas_unidad = (semanas_casos / base_hab) if base_hab > 0 else 0.0
                ind_a = (promedio_semanas_unidad / divisor_periodo) * 100
                ind_b = (u_oportunas / base_hab) * 100
                ind_c = (promedio_semanas_unidad / divisor_periodo) * 100
                ind_d = (u_sin_notificar / base_hab) * 100
                excedente_rsm = max(0.0, ind_d - 5.0)
                ind_e = max(0.0, ind_b - excedente_rsm)
                ind_f = (ind_b + ind_c) / 2.0
                
                processed_results.append({
                    "Unidad": unidad,
                    "a) Cumplimiento u Oportunidad (%)": round(ind_a, 2),
                    "b) Cobertura Oportuna (%)": round(ind_b, 2),
                    "c) Consistencia (%)": round(ind_c, 2),
                    "d) Reporta Sin Movimiento (RSM) (%)": round(ind_d, 2),
                    "e) Cobertura Ajustada (%)": round(ind_e, 2),
                    "f) Calidad (Descriptivo) (%)": round(ind_f, 2),
                    "_raw": {
                        "a": ind_a, "b": ind_b, "c": ind_c, "d": ind_d, "e": ind_e, "f": ind_f
                    },
                    "_metrics": m
                })

            df_resumen = pd.DataFrame(processed_results)

            def get_bg_color(val, ind_type):
                if ind_type == "a":
                    if val == 100.0: return 'background-color: #10B981; color: white; font-weight: bold;'
                    elif 97.5 <= val <= 99.9: return 'background-color: #FFFFFF; color: black; font-weight: bold;'
                    elif 95.0 <= val <= 97.4: return 'background-color: #FEF08A; color: black; font-weight: bold;'
                    else: return 'background-color: #EF4444; color: white; font-weight: bold;'
                elif ind_type in ["b", "e"]:
                    if 95.0 <= val <= 100.0: return 'background-color: #10B981; color: white; font-weight: bold;'
                    elif 90.0 <= val <= 94.9: return 'background-color: #FFFFFF; color: black; font-weight: bold;'
                    elif 80.0 <= val <= 89.9: return 'background-color: #FEF08A; color: black; font-weight: bold;'
                    else: return 'background-color: #EF4444; color: white; font-weight: bold;'
                elif ind_type == "c":
                    if 90.0 <= val <= 100.0: return 'background-color: #10B981; color: white; font-weight: bold;'
                    elif 80.0 <= val <= 89.9: return 'background-color: #FFFFFF; color: black; font-weight: bold;'
                    elif 70.0 <= val <= 79.9: return 'background-color: #FEF08A; color: black; font-weight: bold;'
                    else: return 'background-color: #EF4444; color: white; font-weight: bold;'
                elif ind_type == "d":
                    if 0.0 <= val <= 1.9: return 'background-color: #10B981; color: white; font-weight: bold;'
                    elif 2.0 <= val <= 4.9: return 'background-color: #FFFFFF; color: black; font-weight: bold;'
                    elif 5.0 <= val <= 10.0: return 'background-color: #FEF08A; color: black; font-weight: bold;'
                    else: return 'background-color: #EF4444; color: white; font-weight: bold;'
                elif ind_type == "f":
                    if 90.0 <= val <= 100.0: return 'background-color: #10B981; color: white; font-weight: bold;'
                    elif 80.0 <= val <= 89.9: return 'background-color: #FFFFFF; color: black; font-weight: bold;'
                    elif 60.0 <= val <= 79.9: return 'background-color: #FEF08A; color: black; font-weight: bold;'
                    else: return 'background-color: #EF4444; color: white; font-weight: bold;'
                return ''

            def style_dataframe(row_data):
                styles = [''] * len(row_data)
                col_mapping = {
                    "a) Cumplimiento u Oportunidad (%)": "a",
                    "b) Cobertura Oportuna (%)": "b",
                    "c) Consistencia (%)": "c",
                    "d) Reporta Sin Movimiento (RSM) (%)": "d",
                    "e) Cobertura Ajustada (%)": "e",
                    "f) Calidad (Descriptivo) (%)": "f"
                }
                idx = row_data.name
                raw_dict = df_resumen.loc[idx, "_raw"]
                for i, col_name in enumerate(row_data.index):
                    if col_name in col_mapping:
                        itype = col_mapping[col_name]
                        val = raw_dict[itype]
                        styles[i] = get_bg_color(val, itype)
                return styles

            st.markdown("---")
            st.subheader(f"📊 Tabla General - {trimestre_seleccionado}")
            
            display_df = df_resumen.drop(columns=["_raw", "_metrics"])
            styled_general = display_df.style.format(formatter="{:.2f}", subset=pd.IndexSlice[:, display_df.columns[1:]]).apply(style_dataframe, axis=1)
            st.dataframe(styled_general, use_container_width=True)

            st.markdown("---")
            st.subheader("🏥 Tablas Detalladas por Unidad")
            
            st.markdown("##### 🚦 Leyenda de Acotaciones y Semaforización")
            st.markdown("""
            <div class="legend-container">
                <div class="legend-item legend-excelente">🟢 Excelente</div>
                <div class="legend-item legend-bueno">⚪ Bueno</div>
                <div class="legend-item legend-regular">🟡 Regular</div>
                <div class="legend-item legend-malo">🔴 Malo</div>
            </div>
            """, unsafe_allow_html=True)
            
            unit_options = ["TODAS"] + TARGET_UNITS
            selected_unit = st.selectbox("Seleccione una Unidad Médica (o elija 'TODAS'):", unit_options)
            
            def render_unit_details(unit_name):
                unit_row = df_resumen[df_resumen["Unidad"] == unit_name].iloc[0]
                raw_vals = unit_row["_raw"]
                unit_metrics = unit_row["_metrics"]
                
                st.markdown(f"### 📍 Unidad: **{unit_name}** ({trimestre_seleccionado})")
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    st.markdown("##### Variables Base Mapeadas")
                    var_df = pd.DataFrame(list(unit_metrics.items()), columns=["Variable", "Valor (Columna AB)"])
                    st.dataframe(var_df, use_container_width=True, hide_index=True)
                
                with col2:
                    st.markdown("##### Indicadores y Semáforo")
                    ind_summary = []
                    indicators_meta = [
                        ("a) Cumplimiento u Oportunidad", raw_vals["a"], "a"),
                        ("b) Cobertura Oportuna", raw_vals["b"], "b"),
                        ("c) Consistencia", raw_vals["c"], "c"),
                        ("d) Reporta Sin Movimiento (RSM)", raw_vals["d"], "d"),
                        ("e) Cobertura Ajustada", raw_vals["e"], "e"),
                        ("f) Calidad (Descriptivo)", raw_vals["f"], "f")
                    ]
                    for name, val, itype in indicators_meta:
                        ind_summary.append({
                            "Indicador": name,
                            "Resultado (%)": round(val, 2)
                        })
                    
                    ind_df = pd.DataFrame(ind_summary)
                    
                    def style_ind_table(row_ind):
                        styles = [''] * len(row_ind)
                        itypes = ["a", "b", "c", "d", "e", "f"]
                        idx = row_ind.name
                        itype = itypes[idx] if idx < len(itypes) else "a"
                        for i, col_name in enumerate(row_ind.index):
                            if col_name == "Resultado (%)":
                                val = raw_vals[itype]
                                styles[i] = get_bg_color(val, itype)
                        return styles

                    styled_ind = ind_df.style.format(formatter="{:.2f}", subset=["Resultado (%)"]).apply(style_ind_table, axis=1)
                    st.dataframe(styled_ind, use_container_width=True, hide_index=True)
                st.markdown("---")

            if selected_unit == "TODAS":
                for u in TARGET_UNITS:
                    render_unit_details(u)
            else:
                render_unit_details(selected_unit)

            # -------------------------------------------------------------
            # GENERACIÓN DE REPORTE OFICIAL EN WORD ORGANIZADO POR TRIMESTRES
            # -------------------------------------------------------------
            st.markdown("---")
            st.subheader("📑 Generación de Reporte Oficial en Word (Estructura Trimestral)")
            st.info("Descarga el documento Word completo con las secciones separadas por trimestre (Primer, Segundo, Tercer y Cuarto Trimestre).")

            def get_hex_color(val, ind_type):
                if ind_type == "a":
                    if val == 100.0: return "10B981"
                    elif 97.5 <= val <= 99.9: return "FFFFFF"
                    elif 95.0 <= val <= 97.4: return "FEF08A"
                    else: return "EF4444"
                elif ind_type in ["b", "e"]:
                    if 95.0 <= val <= 100.0: return "10B981"
                    elif 90.0 <= val <= 94.9: return "FFFFFF"
                    elif 80.0 <= val <= 89.9: return "FEF08A"
                    else: return "EF4444"
                elif ind_type == "c":
                    if 90.0 <= val <= 100.0: return "10B981"
                    elif 80.0 <= val <= 89.9: return "FFFFFF"
                    elif 70.0 <= val <= 79.9: return "FEF08A"
                    else: return "EF4444"
                elif ind_type == "d":
                    if 0.0 <= val <= 1.9: return "10B981"
                    elif 2.0 <= val <= 4.9: return "FFFFFF"
                    elif 5.0 <= val <= 10.0: return "FEF08A"
                    else: return "EF4444"
                elif ind_type == "f":
                    if 90.0 <= val <= 100.0: return "10B981"
                    elif 80.0 <= val <= 89.9: return "FFFFFF"
                    elif 60.0 <= val <= 79.9: return "FEF08A"
                    else: return "EF4444"
                return "FFFFFF"

            def generar_reporte_word_trimestral():
                doc = Document()
                for section in doc.sections:
                    section.top_margin = Inches(1)
                    section.bottom_margin = Inches(1)
                    section.left_margin = Inches(1)
                    section.right_margin = Inches(1)
                
                trimestres_lista = ["PRIMER TRIMESTRE", "SEGUNDO TRIMESTRE", "TERCER TRIMESTRE", "CUARTO TRIMESTRE"]
                
                for idx_t, trimestre in enumerate(trimestres_lista):
                    if idx_t > 0:
                        doc.add_page_break()
                        
                    # Encabezado Institucional Oficial por Trimestre
                    p_header = doc.add_paragraph()
                    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_h1 = p_header.add_run("REPRESENTACIÓN REGIONAL SUR\nSUBDELEGACIÓN MÉDICA\nDEPARTAMENTO DE ATENCIÓN MÉDICA\nCOORDINACIÓN DE EPIDEMIOLOGÍA Y MEDICINA PREVENTIVA\n")
                    run_h1.bold = True
                    run_h1.font.size = Pt(8.5)
                    run_h1.font.color.rgb = RGBColor(30, 58, 138)
                    
                    run_h2 = p_header.add_run("INDICADORES PARA EL SISTEMA ÚNICO AUTOMATIZADO DE VIGILANCIA EPIDEMIOLÓGICA (SUAVE)\n")
                    run_h2.bold = True
                    run_h2.font.size = Pt(9.5)
                    
                    run_anio = p_header.add_run(f"AÑO: {anio} | {trimestre} (13 SEMANAS)\n")
                    run_anio.bold = True
                    run_anio.font.size = Pt(9.5)

                    doc.add_paragraph().paragraph_format.space_after = Pt(4)

                    p_title = doc.add_paragraph()
                    r_title = p_title.add_run(f"EVALUACIÓN DE INDICADORES - {trimestre}")
                    r_title.bold = True
                    r_title.font.size = Pt(10)

                    table_gen = doc.add_table(rows=1, cols=5)
                    table_gen.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table_gen.autofit = False

                    headers = [
                        "UNIDAD MÉDICA / TRIMESTRE", 
                        "CUMPLIMIENTO U OPORTUNIDAD", 
                        "COBERTURA OPORTUNA", 
                        "CONSISTENCIA", 
                        "CALIDAD"
                    ]
                    
                    hdr_cells = table_gen.rows[0].cells
                    for i, text in enumerate(headers):
                        hdr_cells[i].text = text
                        p = hdr_cells[i].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(8)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), '1E3A8A')
                        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)

                    itypes_map = ["a", "b", "c", "f"]
                    for idx, row in df_resumen.iterrows():
                        row_cells = table_gen.add_row().cells
                        row_cells[0].text = str(row["Unidad"])
                        raw_dict = row["_raw"]
                        vals = [
                            raw_dict["a"], raw_dict["b"], raw_dict["c"], raw_dict["f"]
                        ]
                        
                        for i, val in enumerate(vals):
                            cell = row_cells[i+1]
                            cell.text = f"{val:.2f}%"
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            
                            hex_c = get_hex_color(val, itypes_map[i])
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), hex_c)
                            cell._tc.get_or_add_tcPr().append(shd)
                            
                            for run in p.runs:
                                run.font.size = Pt(8)
                                if hex_c in ["10B981", "EF4444"]:
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.bold = True
                                else:
                                    run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        p_un = row_cells[0].paragraphs[0]
                        p_un.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in p_un.runs:
                            run.font.size = Pt(8)
                            run.bold = True

                    doc.add_paragraph().paragraph_format.space_after = Pt(12)
                    
                    p_footer = doc.add_paragraph()
                    r_footer = p_footer.add_run(f"Fuente: SINAVE-SUAVE. Cubo de indicadores ({trimestre} {anio}, {total_semanas_anio} semanas totales del año).")
                    r_footer.italic = True
                    r_footer.font.size = Pt(8)
                    r_footer.font.color.rgb = RGBColor(100, 100, 100)

                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                return bio

            word_file = generar_reporte_word_trimestral()
            st.download_button(
                label="📥 Descargar Reporte en Word Organizado por Trimestres (.docx)",
                data=word_file,
                file_name=f"Reporte_Trimestral_Oficial_{anio}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    except Exception as e:
        st.error(f"Ocurrió un error al procesar el archivo Excel: {e}")
else:
    st.info("👈 Por favor, carga tu archivo Excel en la parte superior para comenzar el análisis.")
